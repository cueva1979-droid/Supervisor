import os, re, json
from docx import Document

UPLOADS = r"C:\Users\Analista\Desktop\aplicaciones\PROVEEDORES\Supervisor\backend\uploads"

TITLE_PATTERN = r"(?:Ing\.|Lic\.|Lcda\.|Lcdo\.|Lc\.|Dr\.|Dra\.|Abg\.|Econ\.|Arq\.|Tec\.|Tnlgo\.|Tnlg\.|Tnlga\.|Blga\.|Blg\.|Psic\xf3logo\.|Psic\xf3loga\.|Psic\.|Bi\xf3loga|Biologa|Mgs\.|Msc\.)"

def extract_all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    for section in doc.sections:
        for p in section.header.paragraphs:
            parts.append(p.text)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
    return "\n".join(parts)

def extract_order_number(text):
    m = re.search(r"IC-GADCCC-2026[-]?0*(\d+)", text, re.IGNORECASE)
    if m:
        num = m.group(1)
        return f"IC-GADCCC-2026-{num.zfill(4)}"
    return None

def normalize_title(name):
    """Fix common OCR/encoding issues in names."""
    name = name.strip().rstrip(".")
    # Fix double spaces
    name = re.sub(r"\s+", " ", name)
    # Fix "Lcda." being written as "Lic." or vice versa variants
    return name

def extract_admin_from_header(text):
    """Extract ADMINISTRADOR: from header section."""
    # Build a pattern that captures title + name until newline or certain keywords
    pattern = (
        r"ADMINISTRADOR\s*:\s*"
        + TITLE_PATTERN +
        r"\s*[A-Za-z\u00c1\u00c9\u00cd\u00d3\u00da\u00d1\u00e1\u00e9\u00ed\u00f3\u00fa\u00f1\u00dc\u00fc\s.]+?"
        r"(?=\s*(?:\n|$|\s*(?:N[\.\u00da]?MERO|FISCALIZADOR|PARTIDA|\u00c1REA|CERTIFICACI[O\u00d3]N|COMPROMISO)))"
    )
    m = re.search(pattern, text)
    if m:
        return normalize_title(m.group(0).split(":", 1)[1].strip())
    return None

def extract_admin_from_clause(text):
    """Extract admin from 'La administracion ... estara a cargo de/del ...' clause."""
    pattern = (
        r"La\s+administraci[o\u00f3]n\s+de\s+la\s+orden\s+de\s+compra\s*,\s*"
        r"estar[a\u00e1]\s+a\s+cargo\s+(?:del|de\s+la)\s+"
        + TITLE_PATTERN +
        r"\s*[A-Za-z\u00c1\u00c9\u00cd\u00d3\u00da\u00d1\u00e1\u00e9\u00ed\u00f3\u00fa\u00f1\u00dc\u00fc\s.]+?"
        r"(?:\s*,|\s+quien)"
    )
    m = re.search(pattern, text)
    if m:
        full = m.group(0)
        # Extract just the name part after "cargo del/de la "
        name_part = re.split(r"cargo\s+(?:del|de\s+la)\s+", full, flags=re.IGNORECASE)[-1]
        name_part = re.sub(r"\s*,\s*quien.*$", "", name_part)
        return normalize_title(name_part)
    return None

def extract_admin_combined(text):
    name = extract_admin_from_header(text)
    if name:
        return name
    name = extract_admin_from_clause(text)
    if name:
        return name
    return None

results = []
files = sorted([f for f in os.listdir(UPLOADS) if f.endswith(".docx")])

for fname in files:
    fpath = os.path.join(UPLOADS, fname)
    try:
        doc = Document(fpath)
        text = extract_all_text(doc)
        order_num = extract_order_number(text)
        admin = extract_admin_combined(text)
        results.append({
            "filename": fname,
            "orden": order_num if order_num else "NOT FOUND",
            "administrador": admin if admin else "NOT FOUND"
        })
        status = "OK"
        if not order_num and not admin:
            status = "MISSING BOTH"
        elif not order_num:
            status = "MISSING ORDEN"
        elif not admin:
            status = "MISSING ADMIN"
        print(f"[{status:20s}] {fname}")
        print(f"      Orden: {order_num}")
        print(f"      Admin:  {admin}")
        print()
    except Exception as e:
        results.append({"filename": fname, "orden": "ERROR", "administrador": f"ERROR: {str(e)}"})
        print(f"[ERROR              ] {fname}: {e}")
        print()

print("=" * 80)
print("COMPLETE JSON OUTPUT")
print("=" * 80)
print(json.dumps(results, indent=2, ensure_ascii=False))

print()
print("=" * 135)
print(f'{"#":>3}  {"FILENAME":<67}  {"ORDEN":<28}  {"ADMINISTRADOR":<40}')
print("=" * 135)
for i, r in enumerate(results, 1):
    print(f"{i:>3}  {r['filename']:<67}  {r['orden']:<28}  {r['administrador']:<40}")
print("=" * 135)
