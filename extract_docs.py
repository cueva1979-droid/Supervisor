import os, re, json
from docx import Document

UPLOADS = r"C:\Users\Analista\Desktop\aplicaciones\PROVEEDORES\Supervisor\backend\uploads"

def extract_all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    for section in doc.sections:
        for p in section.header.paragraphs:
            parts.append(p.text)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
    return "\n".join(parts)

def extract_order_number(text):
    m = re.search(r"IC-GADCCC-2026[-]?0*(\d+)", text, re.IGNORECASE)
    if m:
        num = m.group(1)
        return f"IC-GADCCC-2026-{num.zfill(4)}"
    return None

def extract_admin_from_header(text):
    m = re.search(
        r"ADMINISTRADOR\s*:\s*"
        r"((?:Ing\.|Lic\.|Lcda\.|Lcdo\.|Dr\.|Dra\.|Abg\.|Econ\.|Arq\.|Tec\.)\s*"
        r"[A-Za-z\u00c1\u00c9\u00cd\u00d3\u00da\u00d1\u00e1\u00e9\u00ed\u00f3\u00fa\u00f1\u00dc\u00fc\s.]+?)"
        r"(?:\n|$|\s*(?:N.U.MERO|FISCALIZADOR|PARTIDA|.REA|N.MERO|CERTIFICACI.N|COMPROMISO))",
        text
    )
    if m:
        return m.group(1).strip().rstrip(".")
    return None

def extract_admin_from_clause(text):
    m = re.search(
        r"La\s+administraci[o\u00f3]n\s+de\s+la\s+orden\s+de\s+compra\s*,\s*"
        r"estar[a\u00e1]\s+a\s+cargo\s+(?:del|de\s+la)\s+"
        r"((?:Ing\.|Lic\.|Lcda\.|Lcdo\.|Dr\.|Dra\.|Abg\.|Econ\.|Arq\.|Tec\.)\s*"
        r"[A-Za-z\u00c1\u00c9\u00cd\u00d3\u00da\u00d1\u00e1\u00e9\u00ed\u00f3\u00fa\u00f1\u00dc\u00fc\s.]+?)"
        r"(?:\s*,|\s+quien|\s+y\s+como)",
        text
    )
    if m:
        return m.group(1).strip().rstrip(".")
    return None

def extract_admin_combined(text):
    name = extract_admin_from_header(text)
    if name:
        return name
    name = extract_admin_from_clause(text)
    if name:
        return name
    return None

results = []
files = sorted([f for f in os.listdir(UPLOADS) if f.endswith(".docx")])

for fname in files:
    fpath = os.path.join(UPLOADS, fname)
    try:
        doc = Document(fpath)
        text = extract_all_text(doc)
        order_num = extract_order_number(text)
        admin = extract_admin_combined(text)
        results.append({
            "filename": fname,
            "orden": order_num if order_num else "NOT FOUND",
            "administrador": admin if admin else "NOT FOUND"
        })
        status = "OK"
        if not order_num and not admin:
            status = "MISSING BOTH"
        elif not order_num:
            status = "MISSING ORDEN"
        elif not admin:
            status = "MISSING ADMIN"
        print(f"[{status:20s}] {fname}")
        print(f"      Orden: {order_num}")
        print(f"      Admin:  {admin}")
        print()
    except Exception as e:
        results.append({"filename": fname, "orden": "ERROR", "administrador": f"ERROR: {str(e)}"})
        print(f"[ERROR              ] {fname}: {e}")
        print()

print("=" * 80)
print("COMPLETE JSON OUTPUT")
print("=" * 80)
print(json.dumps(results, indent=2, ensure_ascii=False))

print()
print("=" * 130)
print(f'{"#":>3}  {"FILENAME":<65}  {"ORDEN":<28}  {"ADMINISTRADOR":<35}')
print("=" * 130)
for i, r in enumerate(results, 1):
    print(f"{i:>3}  {r['filename']:<65}  {r['orden']:<28}  {r['administrador']:<35}")
print("=" * 130)
