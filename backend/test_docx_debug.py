import sys
sys.path.insert(0, '.')
from docx import Document
import json

doc = Document(r'..\0002.docx')

print("=== PARAGRAPHS ===")
for p in doc.paragraphs:
    if p.text.strip():
        print(repr(p.text.strip()))

print("\n=== TABLES ===")
for ti, table in enumerate(doc.tables):
    print(f"\nTable {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip()[:60] for cell in row.cells]
        print(f"  Row {ri}: {cells}")
        if ri > 30:
            print("  ... (truncated)")
            break
