import re
import pdfplumber
from docx import Document
from typing import List, Dict, Optional, Tuple

MONTHS_ES = {
    'enero': '01', 'febrero': '02', 'marzo': '03', 'abril': '04',
    'mayo': '05', 'junio': '06', 'julio': '07', 'agosto': '08',
    'septiembre': '09', 'octubre': '10', 'noviembre': '11', 'diciembre': '12'
}

class DocumentParser:
    def __init__(self, filepath: str, filename: str):
        self.filepath = filepath
        self.filename = filename
        self.text = ""
        self.tables = []
        self.file_type = "pdf" if filename.lower().endswith('.pdf') else "docx"

    def extract_text(self) -> str:
        if not self.text:
            if self.file_type == "pdf":
                self._extract_pdf()
            else:
                self._extract_docx()
        return self.text

    def _extract_pdf(self):
        try:
            with pdfplumber.open(self.filepath) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        self.text += page_text + "\n"
                    tables = page.extract_tables()
                    if tables:
                        self.tables.extend(tables)
        except Exception as e:
            raise RuntimeError(f"Error al leer PDF: {str(e)}")

    def _extract_docx(self):
        try:
            doc = Document(self.filepath)
            for p in doc.paragraphs:
                if p.text.strip():
                    self.text += p.text.strip() + "\n"
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(cells)
                self.tables.append(rows)
                for row in rows:
                    unique = list(dict.fromkeys(row))
                    joined = " | ".join(unique)
                    if joined.strip():
                        self.text += joined + "\n"
        except Exception as e:
            raise RuntimeError(f"Error al leer DOCX: {str(e)}")

    def _normalize_date(self, date_str: str) -> str:
        date_str = date_str.strip()
        m = re.match(r'(\d{1,2})\s*de\s*([a-záéíóúñ]+)\s*de\s*(\d{4})', date_str, re.IGNORECASE)
        if m:
            return f"{m.group(1).zfill(2)}/{MONTHS_ES.get(m.group(2).lower(), '01')}/{m.group(3)}"
        m = re.match(r'(\d{1,2})[/-](\d{1,2})[/-](\d{2,4})', date_str)
        if m:
            y = m.group(3)
            if len(y) == 2:
                y = '20' + y if int(y) < 50 else '19' + y
            return f"{m.group(1).zfill(2)}/{m.group(2).zfill(2)}/{y}"
        return date_str

    def extract_proveedor(self) -> Optional[str]:
        text = self.extract_text()
        # First: look for PROVEEDOR: in table cells directly (most reliable)
        if self.tables:
            for table in self.tables:
                for row in table:
                    for cell in row:
                        if cell and isinstance(cell, str) and 'PROVEEDOR:' in cell:
                            val = cell.split('PROVEEDOR:')[1].strip().split('\n')[0].strip()
                            if val and 'RUC' not in val.upper() and 'PROFORMA' not in val.upper():
                                return val.split('RUC')[0].strip()
        # Second: regex patterns
        patterns = [
            r'PROVEEDOR:\s*([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑa-záéíóúñ\s]+?)(?:\n|RUC)',
        ]
        for p in patterns:
            m = re.search(p, text, re.IGNORECASE | re.MULTILINE | re.DOTALL)
            if m:
                c = m.group(1).strip().rstrip()
                c = re.sub(r'\s+', ' ', c)
                if len(c) > 3 and 'RUC' not in c.upper() and 'PROFORMA' not in c.upper() and len(c) < 100:
                    return c
        # Third: line by line
        for line in text.split('\n'):
            if 'PROVEEDOR:' in line:
                parts = line.split('PROVEEDOR:')
                if len(parts) > 1:
                    val = parts[1].strip().split('\n')[0].strip()
                    if val:
                        return val.split('RUC')[0].strip()
        return None

    def extract_ruc(self) -> Optional[str]:
        text = self.extract_text()
        m = re.search(r'RUC:\s*(\d{10,20})', text, re.IGNORECASE)
        if m:
            return m.group(1)
        m = re.search(r'RUC\s*[:#]?\s*(\d{6,20})', text, re.IGNORECASE)
        if m:
            return m.group(1)
        m = re.search(r'\b(\d{13})\b', text)
        if m:
            return m.group(1)
        for row in self.tables[0] if self.tables else []:
            for cell in row:
                rm = re.search(r'RUC:\s*(\d{10,20})', cell)
                if rm:
                    return rm.group(1)
        return None

    def extract_administrador(self) -> Optional[str]:
        text = self.extract_text()
        title = r'(?:Ing\.|Arq\.|Arqa\.|Lic\.|Lcda\.|Lcdo\.|Abg\.|Blga\.|Blg\.|Tnlgo\.|Tnlg\.|Tnlga\.|Dr\.|Dra\.|Psic[óo]logo\.|Psic\.|Econ\.|Mgs\.|Msc\.|Mgtr\.|Mtr\.)'
        # First: look in table cells for ADMINISTRADOR: or ADMINISTRADORA:
        if self.tables:
            for table in self.tables:
                for row in table:
                    for cell in row:
                        if cell and isinstance(cell, str):
                            m = re.search(r'ADMINISTRADOR(?:A)?:\s*(' + title + r'\s*[A-ZÁÉÍÓÚÑ][a-záéíóúñ]*(?:\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñ]*)+)', cell, re.IGNORECASE)
                            if m:
                                name = m.group(1).strip()
                                name = re.split(r'\s*(?:FISCALIZADOR|N[UÚ]MERO|PARTIDA|OBJETO)', name)[0].strip()
                                return name
        # Second: regex in full text
        patterns = [
            r'ADMINISTRADOR(?:A)?:\s*(' + title + r'\s*[A-ZÁÉÍÓÚÑ][a-záéíóúñ]*(?:\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñ]*)+)',
            r'designa\s+como\s+administrador\s+(?:al|a\s+la)\s+(' + title + r'\s*[A-ZÁÉÍÓÚÑ][a-záéíóúñ]*(?:\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñ]*)+)',
            r'La\s+administraci[óo]n\s+de\s+la\s+orden\s+de\s+compra[,\s]*est[áa]r[aá]\s+a\s+cargo\s+(?:del|de\s+la)\s+(' + title + r'\s*[A-ZÁÉÍÓÚÑ][a-záéíóúñ]*(?:\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñ]*)+)',
            r'ADMINISTRADOR\s+DE\s+LA\s+ORDEN\s+DE\s+COMPRA\s*\n\s*(' + title + r'\s*[A-ZÁÉÍÓÚÑ][a-záéíóúñ]*(?:\s+[A-ZÁÉÍÓÚÑ][a-záéíóúñ]*)+)',
        ]
        for p in patterns:
            m = re.search(p, text, re.IGNORECASE | re.MULTILINE)
            if m:
                name = m.group(1).strip()
                name = re.split(r'\s*(?:FISCALIZADOR|N[UÚ]MERO|PARTIDA|OBJETO)', name)[0].strip()
                return name
        return None

    def extract_codigo_proceso(self) -> Optional[str]:
        text = self.extract_text()
        patterns = [
            r'(?:C[ÓO]DIGO\s*(?:DEL)?\s*PROCESO|CDP)\s*[:#]?\s*([A-Z0-9\-/]{4,})',
            r'PARTIDA\s*PRESUPUESTARIA\s*[:#]?\s*([\d.]+)',
        ]
        for p in patterns:
            m = re.search(p, text, re.IGNORECASE)
            if m:
                return m.group(1).strip()
        m = re.search(r'\b([A-Z]{2,5}-\d{4}-\d+)\b', text)
        if m:
            return m.group(1)
        # Order number may contain code: IC-GADCCC-2026-0002
        orden = self.extract_numero_orden()
        if orden and '-' in orden:
            parts = orden.split('-')
            if len(parts) >= 3:
                return '-'.join(parts[:3])
        return None

    def extract_numero_orden(self) -> Optional[str]:
        text = self.extract_text()
        # Check table cells first for the most reliable extraction
        if self.tables:
            for table in self.tables:
                for row in table:
                    joined = " ".join(cell for cell in row if cell and isinstance(cell, str))
                    m = re.search(r'(IC-GADCCC-\d{4}-\d+)', joined, re.IGNORECASE)
                    if m:
                        return m.group(1)
        patterns = [
            r'(?:No\.?\s*DE\s*ORDEN\s*DE\s*COMPRA|N[°º]\s*DE\s*ORDEN\s*DE\s*COMPRA)\s*(?:\|[^|]+\|)?\s*([A-Z]{2,5}-[A-Z0-9]+-\d{4}-\d+)',
            r'(IC-GADCCC-\d{4}-\d+)',
            r'(?:ORDEN\s*(?:DE\s*)?(?:COMPRA)?\s*[:#]?\s*|OC\s*[:#]?\s*)([\w\-/]{5,})',
        ]
        for p in patterns:
            m = re.search(p, text, re.IGNORECASE)
            if m:
                return m.group(1).strip()
        return None

    def extract_fecha(self) -> Optional[str]:
        text = self.extract_text()
        m = re.search(r'FECHA:\s*(.+?)(?:\n|$)', text, re.IGNORECASE)
        if m:
            try:
                return self._normalize_date(m.group(1).strip().rstrip('.'))
            except:
                pass
        m = re.search(r'(\d{1,2})\s*de\s*([a-záéíóúñ]+)\s*de\s*(\d{4})', text, re.IGNORECASE)
        if m:
            return f"{m.group(1).zfill(2)}/{MONTHS_ES.get(m.group(2).lower(), '01')}/{m.group(3)}"
        m = re.search(r'\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b', text)
        if m:
            return self._normalize_date(m.group(1))
        from datetime import datetime
        return datetime.now().strftime("%d/%m/%Y")

    def extract_objeto(self) -> Optional[str]:
        text = self.extract_text()
        patterns = [
            r'OBJETO\s*DE\s*CONTRATACI[OÓ]N\s*[:#]?\s*["""\u201c\u201d]*(.+?)["""\u201c\u201d]*(?:\n[A-Z\s]{2,}|\Z)',
            r'OBJETO\s*DE\s*CONTRATACI[OÓ]N\s*[:#]?\s*(.+)$',
        ]
        for p in patterns:
            m = re.search(p, text, re.IGNORECASE | re.DOTALL | re.MULTILINE)
            if m:
                r = m.group(1).strip().strip('"').strip('\u201c\u201d')
                if len(r) > 10:
                    return r[:300]
        return None

    def extract_items(self) -> List[Dict]:
        items = []
        # Ensure tables are extracted (calls extract_text which populates self.tables)
        text = self.extract_text()
        # Strategy 1: parse item tables from structured table data
        if self.tables:
            for table in self.tables:
                parsed = self._parse_item_table_v2(table)
                if parsed:
                    items.extend(parsed)
        # Strategy 2: fallback to text-based extraction
        if not items:
            items = self._parse_items_from_text(text)
        # Strategy 3: line-by-line fallback
        if not items:
            items = self._parse_items_from_lines(text)
        # Default single item if nothing found
        if not items:
            items.append({
                "codigo_cpc": "requiere revisión",
                "descripcion": "requiere revisión",
                "cantidad": 1.0, "unidad": "requiere revisión",
                "precio_unitario": 0.0, "subtotal": 0.0, "requires_review": True
            })
        for item in items:
            item["subtotal"] = round(item.get("cantidad", 1) * item.get("precio_unitario", 0), 2)
        return items

    def _parse_item_table_v2(self, table: List[List[str]]) -> List[Dict]:
        """Robust item table parser using header detection + content heuristics."""
        if not table or len(table) < 3:
            return []
        CPC_RE = re.compile(r'^\d{8,}$')
        ncols = max(len(r) for r in table)
        if ncols == 0:
            return []
        cpc_col = desc_col = qty_col = price_col = None

        # Count rows whose first field is a digit (potential item rows)
        data_row_count = sum(1 for r in table if r and str(r[0] or '').strip().isdigit())

        def norm_hdr(s):
            return s.strip().upper().replace(' ', '').replace('.', '').replace(',', '').replace('Ó', 'O').replace('Í', 'I').replace('É', 'E').replace('Á', 'A')

        # Step 1: Header detection
        header_kw = {
            'cpc': ['CPC', 'CODIGO', 'COD', 'CÓDIGO', 'CÓD'],
            'desc': ['DESCRIPCION', 'DESCRIPCIÓN', 'DESCRIP', 'DETALLE', 'DETALLES', 'RUBRO', 'CONCEPTO'],
            'qty': ['CANT', 'CANTIDAD', 'CANT.', 'CANTIDADES'],
            'price': ['UNITARIO', 'V.UNIT', 'V.UNITARIO', 'V. UNITARIO', 'PRECIO', 'VALOR UNITARIO'],
        }
        header_row_idx = None
        for idx, row in enumerate(table):
            row_text = ' '.join(str(c or '').strip() for c in row).upper()
            found_cats = set()
            for cat in ['cpc', 'desc', 'qty', 'price']:
                for kw in header_kw[cat]:
                    if kw in row_text:
                        found_cats.add(cat)
                        break
            if len(found_cats) >= 3:
                header_row_idx = idx
                break

        if header_row_idx is not None:
            hdr = table[header_row_idx]
            for i in range(min(len(hdr), ncols)):
                cell = norm_hdr(str(hdr[i] or ''))
                if not cell: continue
                if cell == 'CPC' or ('CPC' in cell and cell not in ('NRO', 'ITEM', '#')):
                    cpc_col = i
                if any(kw in cell for kw in ['DESCRIPCION', 'DESCRIP', 'DETALLE', 'DETALLES', 'RUBRO', 'CONCEPTO']):
                    desc_col = i
                if any(kw in cell for kw in ['CANT', 'CANTIDAD']):
                    qty_col = i
                if any(kw in cell for kw in ['UNITARIO', 'PRECIO', 'VUNIT']):
                    price_col = i

            # Validate CPC column: if header-detected column doesn't have CPC data, use fallback
            if cpc_col is not None:
                data_rows = [r for r in table if r and len(r) > cpc_col and str(r[0] or '').strip().isdigit()]
                if data_rows:
                    cpc_hits = sum(1 for r in data_rows if CPC_RE.match(str(r[cpc_col] or '').strip()))
                    if cpc_hits / len(data_rows) < 0.5:
                        cpc_col = None

        # Step 2: Column statistics
        col_stats = {}
        for i in range(ncols):
            col_stats[i] = {"nums": [], "decimals": 0, "texts": [], "cpc_count": 0, "unique_texts": set()}
        for row in table:
            for i in range(min(len(row), ncols)):
                val = str(row[i] or "").strip()
                if not val: continue
                if CPC_RE.match(val):
                    col_stats[i]["cpc_count"] += 1
                num = self._parse_number(val)
                if num > 0:
                    col_stats[i]["nums"].append(num)
                    if num != int(num):
                        col_stats[i]["decimals"] += 1
                col_stats[i]["unique_texts"].add(val)
                if len(val) > 5:
                    col_stats[i]["texts"].append(val)

        # Step 3: CPC fallback
        if cpc_col is None:
            best_cpc = 0
            for i, st in col_stats.items():
                if st["cpc_count"] > best_cpc:
                    best_cpc = st["cpc_count"]
                    cpc_col = i
            if best_cpc == 0:
                cpc_col = None
        if cpc_col is None:
            return []

        # Step 4: Item-number column detection (usually col 0)
        item_num_col = None
        st0 = col_stats.get(0)
        if st0 and len(st0["nums"]) >= 3:
            sorted_nums = sorted(set(st0["nums"]))
            if len(sorted_nums) >= 3 and sorted_nums[:3] == [1.0, 2.0, 3.0]:
                item_num_col = 0

        def is_dup(i):
            if i == 0: return False
            common = len(col_stats[i]["unique_texts"] & col_stats[i-1]["unique_texts"])
            total = len(col_stats[i]["unique_texts"]) + len(col_stats[i-1]["unique_texts"])
            if total == 0: return False
            return common / max(total, 1) > 0.8

        def is_unit_col(i):
            st = col_stats[i]
            uv = st["unique_texts"]
            if not uv: return False
            non_numeric = sum(1 for v in uv if self._parse_number(v) == 0 and v.strip() != '')
            return non_numeric >= 2 and non_numeric / len(uv) > 0.5

        min_thresh = min(3, max(1, data_row_count))
        # Step 5: Description column
        if desc_col is None:
            for i in range(cpc_col + 1, ncols):
                if len(col_stats[i]["texts"]) >= min_thresh and not is_dup(i):
                    desc_col = i
                    break
            if desc_col is None:
                for i in range(cpc_col + 1, ncols):
                    if len(col_stats[i]["texts"]) >= min_thresh:
                        desc_col = i
                        break
        if desc_col is None:
            return []

        # Step 6: Quantity column
        if qty_col is None:
            for i in range(desc_col + 1, ncols):
                if is_dup(i) or is_unit_col(i) or i == cpc_col or i == item_num_col:
                    continue
                small = [n for n in col_stats[i]["nums"] if 0 < n < 100000]
                if len(small) >= min_thresh:
                    qty_col = i
                    break
        if qty_col is None:
            for i in range(ncols - 1, desc_col, -1):
                if is_dup(i) or is_unit_col(i) or i == cpc_col or i == item_num_col:
                    continue
                small = [n for n in col_stats[i]["nums"] if 0 < n < 100000]
                if len(small) >= min_thresh:
                    qty_col = i
                    break

        # Step 7: Price column
        if price_col is None and qty_col is not None:
            for i in range(qty_col + 1, ncols):
                if is_dup(i) or i == cpc_col: continue
                if len([n for n in col_stats[i]["nums"] if n >= 0.01]) >= min_thresh:
                    price_col = i
                    break
        if price_col is None and qty_col is not None:
            for i in range(ncols - 1, qty_col, -1):
                if is_dup(i) or i == cpc_col: continue
                if len([n for n in col_stats[i]["nums"] if n >= 0.01]) >= min_thresh:
                    price_col = i
                    break

        # Step 8: Verify against item-number column
        if qty_col is not None and item_num_col is not None:
            data_rows = [r for r in table if len(r) > max(qty_col, desc_col, cpc_col) and
                         str(r[0] or '').strip().isdigit()]
            if data_rows:
                qty_samples = [self._parse_number(str(r[qty_col] or '')) for r in data_rows[:5] if qty_col < len(r)]
                item_samples = [self._parse_number(str(r[item_num_col] or '')) for r in data_rows[:5] if item_num_col < len(r)]
                matches = sum(1 for q, n in zip(qty_samples, item_samples) if 0 < q == n < 100)
                if matches >= min_thresh:
                    qty_col = None
                    for i in range(desc_col + 1, ncols):
                        if is_dup(i) or is_unit_col(i) or i == cpc_col or i == item_num_col:
                            continue
                        small = [n for n in col_stats[i]["nums"] if 0 < n < 1000]
                        if len(small) >= 3:
                            qty_col = i
                            break
                    if qty_col is not None:
                        price_col = None
                        for i in range(qty_col + 1, ncols):
                            if is_dup(i) or i == cpc_col: continue
                            if len([n for n in col_stats[i]["nums"] if n >= 0.01]) >= 3:
                                price_col = i
                                break

        # Step 9: Swap detection
        if qty_col is not None and price_col is not None:
            data_rows = [r for r in table if len(r) > max(qty_col, price_col, desc_col, cpc_col) and
                         str(r[0] or '').strip().isdigit()]
            if data_rows:
                qty_s = [self._parse_number(str(r[qty_col] or '')) for r in data_rows[:5] if qty_col < len(r)]
                price_s = [self._parse_number(str(r[price_col] or '')) for r in data_rows[:5] if price_col < len(r)]
                valid = [(q, p) for q, p in zip(qty_s, price_s) if q > 0 and p > 0]
                if len(valid) >= 2:
                    large_qty = sum(1 for q, _ in valid if q > 100)
                    small_price = sum(1 for _, p in valid if p < 100)
                    if large_qty >= min_thresh and small_price >= min_thresh:
                        qty_col, price_col = price_col, qty_col

        # Step 10: Parse item rows
        items = []
        for row in table:
            first_val = str(row[0] or "").strip()
            if not first_val: continue
            if first_val.upper() in ["SUBTOTAL", "TOTAL", "NOTAS:", "NRO.", "NRO", "ITEM", "SUB TOTAL"]:
                continue
            if not first_val.isdigit(): continue
            cpc_val = str(row[cpc_col] or "").strip() if cpc_col < len(row) else ""
            if cpc_val and not CPC_RE.match(cpc_val): continue
            desc_val = str(row[desc_col] or "").strip() if desc_col < len(row) else ""
            if not desc_val or len(desc_val) < 5: continue
            qty = self._parse_number(str(row[qty_col] or "")) if qty_col is not None and qty_col < len(row) else 1.0
            price = self._parse_number(str(row[price_col] or "")) if price_col is not None and price_col < len(row) else 0.0
            if qty == 0: qty = 1.0
            # Extract unit
            unit_val = ""
            for ui in range(desc_col + 1, ncols):
                if ui in (qty_col, price_col, cpc_col) or is_dup(ui): continue
                if is_unit_col(ui) and ui < len(row):
                    uv = str(row[ui] or '').strip()
                    if uv and len(uv) <= 10 and self._parse_number(uv) == 0:
                        unit_val = uv
                        break
            items.append({
                "codigo_cpc": cpc_val,
                "descripcion": desc_val,
                "cantidad": qty,
                "unidad": unit_val,
                "precio_unitario": price,
                "subtotal": 0.0,
                "requires_review": False
            })
        return items

    def _parse_items_from_text(self, text: str) -> List[Dict]:
        items = []
        lines = text.split('\n')
        in_items = False
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue
            if re.search(r'(?:item|ítem|cantidad|código\s*cpc|descripcion|rubro)', stripped, re.IGNORECASE):
                in_items = True
                continue
            if in_items and re.search(r'(?:subtotal|total\s*general|observacion|nota|forma\s*de\s*pago|plazo|garant)', stripped, re.IGNORECASE):
                break
            if not in_items:
                continue
            parts = re.split(r'\s{2,}|\t', stripped)
            parts = [p.strip() for p in parts if p.strip()]
            if len(parts) >= 3:
                qty = self._parse_number(parts[0])
                desc = parts[1]
                precio = self._parse_number(parts[-1])
                if qty > 0 and len(desc) > 3:
                    items.append({
                        "codigo_cpc": "", "descripcion": desc,
                        "cantidad": qty, "unidad": "",
                        "precio_unitario": precio, "subtotal": 0.0, "requires_review": False
                    })
        return items

    def _parse_items_from_lines(self, text: str) -> List[Dict]:
        items = []
        pat = re.compile(r'(\d[\d.,]*)\s+([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑa-záéíóúñ\s,]+?)\s+(\d[\d.,]*(?:\s*\d[\d.,]*)?)\s*$')
        for line in text.split('\n'):
            s = line.strip()
            if not s or len(s) < 10:
                continue
            if re.search(r'(?:cpc|código|descripcion|total|subtotal|firma|ingeniero|director|administrador|forma|plazo|garant|multa|lugar|recepción)', s, re.IGNORECASE):
                continue
            m = pat.match(s)
            if m:
                qty = self._parse_number(m.group(1))
                desc = m.group(2).strip()
                precio = self._parse_number(m.group(3))
                if qty > 0 and len(desc) > 3:
                    items.append({
                        "codigo_cpc": "", "descripcion": desc,
                        "cantidad": qty, "unidad": "",
                        "precio_unitario": precio, "subtotal": 0.0, "requires_review": False
                    })
        return items

    def _parse_number(self, value: Optional[str]) -> float:
        if not value or not value.strip():
            return 0.0
        if re.search(r'[A-Za-z\u00C0-\u024F]', value.strip()):
            return 0.0
        cleaned = re.sub(r'[^\d.,]', '', value.strip())
        if ',' in cleaned and '.' in cleaned:
            if cleaned.rindex(',') > cleaned.rindex('.'):
                cleaned = cleaned.replace('.', '').replace(',', '.')
            else:
                cleaned = cleaned.replace(',', '')
        elif ',' in cleaned:
            cleaned = cleaned.replace(',', '.')
        try:
            return float(cleaned)
        except ValueError:
            return 0.0

    def extract_plazo_entrega(self) -> Optional[str]:
        text = self.extract_text()
        patterns = [
            r'PLAZO\s*DE\s*ENTREGA\s*[:#]?\s*(.+?)(?:\n[A-Z\s]{3,}|$)',
            r'PLAZO\s*[:#]?\s*(?:DE\s*ENTREGA\s*)?[:#]?\s*(.+?)(?:\n[A-Z\s]{3,}|$)',
        ]
        for p in patterns:
            m = re.search(p, text, re.IGNORECASE | re.DOTALL | re.MULTILINE)
            if m:
                r = m.group(1).strip().strip('"').strip('\u201c\u201d').rstrip('.')
                if len(r) > 2:
                    return r[:100]
        for line in text.split('\n'):
            if 'PLAZO' in line.upper() and ('ENTREGA' in line.upper() or 'ENTREGA' not in text.upper()):
                m = re.search(r'PLAZO(?:\s*DE\s*ENTREGA)?\s*[:#]?\s*(.+)', line, re.IGNORECASE)
                if m:
                    r = m.group(1).strip().rstrip('.')
                    if len(r) > 2:
                        return r[:100]
        return None

    def extract_monto_total(self, items: List[Dict]) -> float:
        total = sum(item.get("subtotal", 0) for item in items)
        if total > 0:
            return round(total, 2)
        text = self.extract_text()
        m = re.search(r'SUBTOTAL[^0-9]*(\d[\d,.]*\d)', text, re.IGNORECASE)
        if m:
            return self._parse_number(m.group(1))
        for p in [r'(?:TOTAL|MONTO\s*TOTAL|IMPORTE\s*TOTAL)\s*[:#]?\s*[\s$]*(\d[\d,.]*\d)',
                  r'SUBTOTAL\s+(?:\w+\s+){0,5}(\d[\d,.]*\d)']:
            m = re.search(p, text, re.IGNORECASE)
            if m:
                return self._parse_number(m.group(1))
        for row in text.split('\n'):
            if 'subtotal' in row.lower():
                nums = re.findall(r'(\d[\d,.]*\d)', row)
                if nums:
                    return self._parse_number(nums[-1])
        return 0.0

    def get_all_data(self) -> Dict:
        items = self.extract_items()
        return {
            "proveedor": self.extract_proveedor() or "requiere revisión",
            "ruc": self.extract_ruc() or "requiere revisión",
            "codigo_proceso": self.extract_codigo_proceso() or "requiere revisión",
            "numero_orden": self.extract_numero_orden() or "requiere revisión",
            "fecha": self.extract_fecha() or "requiere revisión",
            "objeto_contratacion": self.extract_objeto() or "requiere revisión",
            "administrador": self.extract_administrador(),
            "plazo_entrega": self.extract_plazo_entrega(),
            "items": items,
            "monto_total": self.extract_monto_total(items),
        }
