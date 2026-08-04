import io
from typing import Optional


def format_currency(value) -> str:
    if value is None:
        return "$0.00"
    try:
        return f"${float(value):,.2f}"
    except (ValueError, TypeError):
        return "$0.00"


def generate_certificate_docx(data: dict) -> Optional[bytes]:
    """Generate a .docx certificate from form data using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        raise Exception("python-docx is required for certificate generation")

    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"CERTIFICACIÓN PAC NRO. {data.get('cert_nro', 'S/N')}")
    run.bold = True
    run.font.size = Pt(14)
    run.underline = True

    doc.add_paragraph()

    # Intro
    p = doc.add_paragraph()
    p.add_run("Quien suscribe, previo a iniciar el procedimiento denominado: ").bold = True
    p.add_run(f"{data.get('objeto', 'No especificado')}.")

    doc.add_paragraph()

    # Certifico
    p = doc.add_paragraph()
    p.add_run("CERTIFICO: ").bold = True
    p.add_run(data.get("base_legal", ""))

    doc.add_paragraph()

    # Data table
    table_data = [
        ("PARTIDA PRESUPUESTARIA:", data.get("partida", "")),
        ("CPC:", data.get("cpc", "")),
        ("TIPO DE COMPRA:", data.get("tipo_compra", "")),
        ("TIPO DE RÉGIMEN:", data.get("tipo_regimen", "")),
        ("PROCEDIMIENTO:", data.get("procedimiento", "")),
        ("DETALLE:", data.get("objeto", "")),
        ("VALOR ESTIMADO ($):", format_currency(data.get("valor", 0))),
        ("CUATRIMESTRE DE EJECUCIÓN:", data.get("periodo", "")),
        ("CONSTA PAC:", data.get("verificacion_catalogo", "SI")),
    ]

    table = doc.add_table(rows=len(table_data), cols=2)
    table.style = 'Table Grid'
    for i, (key, value) in enumerate(table_data):
        cell_0 = table.cell(i, 0)
        cell_1 = table.cell(i, 1)
        cell_0.text = key
        cell_1.text = str(value)
        for paragraph in cell_0.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    doc.add_paragraph()

    # Signature table
    sig_table = doc.add_table(rows=1, cols=3)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Elaborado por
    cell = sig_table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ELABORADO POR")
    run.bold = True
    cell.paragraphs[0].add_run(f"\n{data.get('elaborado_por', '')}")
    cell.paragraphs[0].add_run(f"\n{data.get('cargo', '')}")

    # Empty middle cell
    sig_table.cell(0, 1).text = ""

    # Aprobado por
    cell = sig_table.cell(0, 2)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("APROBADO POR")
    run.bold = True
    cell.paragraphs[0].add_run(f"\n{data.get('aprobado_por', '')}")
    cell.paragraphs[0].add_run(f"\n{data.get('cargo_aprobado', '')}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def generate_certificate_from_template(doc_data: dict, template_path: str) -> Optional[bytes]:
    """Generate .docx certificate from a template file using docxtemplater-style approach."""
    try:
        from docx import Document
    except ImportError:
        raise Exception("python-docx is required")

    doc = Document(template_path)

    # Replace placeholders in paragraphs
    replacements = {
        "{cert_nro}": "S/N",
        "{anio}": str(doc_data.get("anio", "")),
        "{objeto}": doc_data.get("descripcion", ""),
        "{base_legal}": "Que conforme a lo establecido en el Art.-66.- Reglamento de la Ley Orgánica del Sistema de Contratación Pública-LOSNCP.",
        "{partida}": doc_data.get("partida_presupuestaria", ""),
        "{cpc}": doc_data.get("cpc", ""),
        "{tipo_compra}": doc_data.get("tipo_compra", ""),
        "{tipo_regimen}": doc_data.get("tipo_regimen", ""),
        "{procedimiento}": doc_data.get("procedimiento", ""),
        "{descripcion}": doc_data.get("descripcion", ""),
        "{costo_unitario}": format_currency(doc_data.get("costo_unitario", 0)),
        "{valor}": format_currency(doc_data.get("costo_unitario", 0)),
        "{periodo}": doc_data.get("periodo", ""),
        "{lugar}": "Quito",
        "{fecha}": "",
        "{fecha_actual}": "",
    }

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for key, value in replacements.items():
                if key in run.text:
                    run.text = run.text.replace(key, value)

    # Also replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for key, value in replacements.items():
                            if key in run.text:
                                run.text = run.text.replace(key, value)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
